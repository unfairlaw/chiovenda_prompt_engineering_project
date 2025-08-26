from io import BytesIO
import requests
import pypdf
from docx import Document
import re
import sys
import os
import glob
from typing import List, Set, Optional

def clean_text_enhanced(text: str, 
                       repeated_expressions: Optional[Set[str]] = None,
                       min_word_threshold: int = 3) -> List[str]:
    """
    Enhanced text cleaning function that:
    1. Preserves text alignment and paragraph structure
    2. Removes repeated expressions found across document pages
    3. Removes specific legal document footers
    4. Maintains readability while cleaning whitespace
    """
    if not text:
        return []
    
    # Remove specific legal document expressions
    text = remove_legal_footers(text)
    
    # Split into lines and preserve structure
    lines = text.split('\n')
    
    # Auto-detect repeated expressions if not provided
    if repeated_expressions is None:
        repeated_expressions = detect_repeated_expressions(lines)
    
    # Clean lines while preserving alignment
    cleaned_lines = []
    for line in lines:
        cleaned_line = clean_line_preserve_alignment(line, repeated_expressions, min_word_threshold)
        if cleaned_line is not None:
            cleaned_lines.append(cleaned_line)
    
    # Group into paragraphs
    paragraphs = group_into_paragraphs(cleaned_lines)
    
    return paragraphs

def remove_legal_footers(text: str) -> str:
    """Remove specific legal document footers and expressions"""
    patterns_to_remove = [
        # Original access link pattern
        r'Para conferir o original, acesse o site https://esaj\.tjsp\.jus\.br/[^\n]*',
        # Digital signature pattern
        r'Este documento é cópia do original, assinado digitalmente por [^\n]*',
        # Common variations
        r'Para conferir.*?https://esaj\.tjsp\.jus\.br[^\n]*',
        r'Este documento.*?assinado digitalmente.*?[\n\r]',
        # Page numbers and document metadata
        r'Página \d+ de \d+',
        r'Processo n[°º]?\s*\d+[\d\.\-/]*',
        # Common legal document headers
        r'^TRIBUNAL.*[\n\r]',
        r'^PODER JUDICIÁRIO.*[\n\r]',
        r'^COMARCA DE.*[\n\r]',
        r'^FORO.*[\n\r]',
        r'DOCUMENTO ASSINADO DIGITALMENTE NOS TERMOS DA LEI.*[\n\r]',
    ]
    
    for pattern in patterns_to_remove:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
    
    return text

def detect_repeated_expressions(lines: List[str], threshold: int = 3) -> Set[str]:
    """Auto-detect expressions that appear repeatedly across the document"""
    line_counts = {}
    repeated_expressions = set()
    
    for line in lines:
        clean_line = line.strip()
        if len(clean_line) > 10:  # Only consider substantial lines
            line_counts[clean_line] = line_counts.get(clean_line, 0) + 1
    
    # Find lines that appear multiple times
    for line, count in line_counts.items():
        if count >= threshold:
            repeated_expressions.add(line)
    
    return repeated_expressions

def clean_line_preserve_alignment(line: str, repeated_expressions: Set[str], min_word_threshold: int) -> Optional[str]:
    """Clean individual line while preserving meaningful alignment"""
    stripped_line = line.strip()
    
    # Skip empty lines
    if not stripped_line:
        return ""
    
    # Remove repeated expressions
    if stripped_line in repeated_expressions:
        return None
    
    # Remove lines with too few words (likely artifacts)
    if len(stripped_line.split()) < min_word_threshold:
        # Exception for numbered items, dates, or legal references
        if not re.match(r'^\d+[\.\)]\s*|^\d{1,2}/\d{1,2}/\d{4}|^Art\.?\s*\d+', stripped_line):
            return None
    
    # Preserve indentation for structured content
    leading_spaces = len(line) - len(line.lstrip())
    if leading_spaces > 0 and leading_spaces <= 10:
        preserved_indent = "  " * min(leading_spaces // 4, 2)
        return preserved_indent + stripped_line
    
    return stripped_line

def group_into_paragraphs(lines: List[str]) -> List[str]:
    """Group cleaned lines into coherent paragraphs while preserving structure"""
    paragraphs = []
    current_paragraph = []
    
    for line in lines:
        if line == "":  # Empty line - paragraph break
            if current_paragraph:
                paragraphs.append(" ".join(current_paragraph))
                current_paragraph = []
        elif line.startswith("  "):  # Indented line - likely new item/section
            if current_paragraph:
                paragraphs.append(" ".join(current_paragraph))
                current_paragraph = []
            current_paragraph.append(line.strip())
        else:
            current_paragraph.append(line)
    
    # Add final paragraph
    if current_paragraph:
        paragraphs.append(" ".join(current_paragraph))
    
    # Remove very short paragraphs that are likely artifacts
    return [p for p in paragraphs if len(p.split()) >= 3 or re.match(r'^\d+[\.\)]|^Art\.', p)]

def process_pdf(pdf_stream, output_filename):
    """Process a PDF stream and save as DOCX with enhanced text cleaning"""
    try:
        # Create a PdfReader object to read the PDF
        reader = pypdf.PdfReader(pdf_stream)
        
        # Create a new Document object (for .docx)
        doc = Document()
        
        # Collect all text from all pages for better repeated expression detection
        all_pages_text = []
        page_texts = []
        
        # First pass: extract all text
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text and text.strip():
                    page_texts.append(text)
                    all_pages_text.append(text)
                else:
                    page_texts.append("")
                    print(f"  Warning: Page {i + 1} appears to be empty or contains no extractable text")
            except Exception as e:
                print(f"  Error extracting text from page {i + 1}: {e}")
                page_texts.append("")
        
        # Detect repeated expressions across all pages
        combined_text = "\n".join(all_pages_text)
        combined_lines = combined_text.split('\n')
        repeated_expressions = detect_repeated_expressions(combined_lines, threshold=2)
        
        print(f"  Detected {len(repeated_expressions)} repeated expressions to remove")
        
        # Second pass: process each page with enhanced cleaning
        processed_pages = 0
        total_paragraphs = 0
        
        for i, text in enumerate(page_texts):
            if not text:
                continue
                
            try:
                # Clean the text using enhanced function
                paragraphs = clean_text_enhanced(text, repeated_expressions)
                
                if paragraphs:
                    # Add page header only if there's content
                    doc.add_heading(f'Page {i + 1}', level=2)
                    
                    # Add each paragraph separately
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            doc.add_paragraph(paragraph)
                    
                    processed_pages += 1
                    total_paragraphs += len(paragraphs)
                    print(f"  Processed page {i + 1}: {len(paragraphs)} paragraphs extracted")
                else:
                    print(f"  Page {i + 1}: No content after cleaning")
                
            except Exception as e:
                print(f"  Error processing page {i + 1}: {e}")
                continue
        
        if processed_pages == 0:
            print("  Warning: No content was extracted from any page")
            return False
        
        # Save the document as a .docx file
        doc.save(output_filename)
        print(f"Text successfully extracted and saved to '{output_filename}'")
        print(f"Summary: {processed_pages} pages processed, {total_paragraphs} paragraphs extracted")
        return True
        
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return False

def process_directory(directory_path):
    """Process all PDF files in a directory"""
    # Find all PDF files in the directory
    pdf_pattern = os.path.join(directory_path, "*.pdf")
    pdf_files = glob.glob(pdf_pattern)
    
    # Also check for PDF files with uppercase extension
    pdf_pattern_upper = os.path.join(directory_path, "*.PDF")
    pdf_files.extend(glob.glob(pdf_pattern_upper))
    
    if not pdf_files:
        print(f"No PDF files found in directory: {directory_path}")
        return
    
    print(f"Found {len(pdf_files)} PDF file(s) in directory: {directory_path}")
    
    successful_conversions = 0
    failed_conversions = 0
    
    for pdf_file in pdf_files:
        print(f"\nProcessing: {os.path.basename(pdf_file)}")
        
        try:
            with open(pdf_file, 'rb') as file:
                pdf_stream = BytesIO(file.read())
                base_name = os.path.splitext(os.path.basename(pdf_file))[0]
                output_filename = os.path.join(directory_path, f"{base_name}_extracted.docx")
                
                if process_pdf(pdf_stream, output_filename):
                    successful_conversions += 1
                else:
                    failed_conversions += 1
                    
        except Exception as e:
            print(f"Error reading PDF file {pdf_file}: {e}")
            failed_conversions += 1
    
    print(f"\n--- Conversion Summary ---")
    print(f"Successful conversions: {successful_conversions}")
    print(f"Failed conversions: {failed_conversions}")
    print(f"Total files processed: {len(pdf_files)}")

def main():
    if len(sys.argv) != 2:
        print("""Enhanced PDF to DOCX Text Extractor with Advanced Cleaning
        
Usage: 
python pdf_docx_extractor.py "<pdf_url_or_file_path_or_directory>"
        
Features:
- Removes repeated headers/footers across pages
- Removes legal document footers (ESAJ links, digital signatures)
- Preserves text alignment and structure
- Handles single files, directories, or URLs
        
Examples:
- Single file: python pdf_docx_extractor.py "document.pdf"
- Directory: python pdf_docx_extractor.py "/path/to/pdf/directory"
- URL: python pdf_docx_extractor.py "https://example.com/document.pdf" """)
        sys.exit(1)
    
    pdf_source = sys.argv[1]
    
    # Check if source exists
    if not pdf_source.startswith("http") and not os.path.exists(pdf_source):
        print(f"Error: Path does not exist: {pdf_source}")
        sys.exit(1)
    
    if pdf_source.startswith("http"):
        # Handle URL
        try:
            print(f"Downloading PDF from {pdf_source}...")
            response = requests.get(pdf_source, timeout=30)
            if response.status_code != 200:
                print(f"Failed to download PDF from {pdf_source}. Status code: {response.status_code}")
                sys.exit(1)
            
            pdf_stream = BytesIO(response.content)
            output_filename = "extracted_text_from_url.docx"
            process_pdf(pdf_stream, output_filename)
            
        except Exception as e:
            print(f"Error downloading PDF: {e}")
            sys.exit(1)
    
    elif os.path.isdir(pdf_source):
        # Handle directory
        process_directory(pdf_source)
    
    elif os.path.isfile(pdf_source):
        # Handle single local file
        if not pdf_source.lower().endswith('.pdf'):
            print(f"Error: File is not a PDF: {pdf_source}")
            sys.exit(1)
            
        try:
            print(f"Processing single PDF file: {os.path.basename(pdf_source)}")
            with open(pdf_source, 'rb') as pdf_file:
                pdf_stream = BytesIO(pdf_file.read())
                base_name = os.path.splitext(os.path.basename(pdf_source))[0]
                output_filename = f"{base_name}_extracted.docx"
                process_pdf(pdf_stream, output_filename)
                
        except Exception as e:
            print(f"Error reading local PDF file: {e}")
            sys.exit(1)
    
    else:
        print(f"Invalid source: {pdf_source}")
        print("Please provide a valid URL, file path, or directory path.")
        sys.exit(1)

if __name__ == "__main__":
    main()